"""
Generate Panduan Penggunaan SIDATA - Sistem Data ASN SILAKAP V1.0
Output: SIDATA_Panduan_Penggunaan.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'CCCCCC')
                tcBorders.append(border)
            tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x0D, 0x47, 0x2E)
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x1A, 0x6B, 0x47)
        run.font.size = Pt(13)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x2D, 0x7A, 0x55)
        run.font.size = Pt(11)
    return p

def add_para(doc, text, bold=False, italic=False, indent=False, size=10.5, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ' ')
        r1.font.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_info_box(doc, title, content_lines, color_hex='EAF4EE'):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, color_hex)
    cell.paragraphs[0].clear()
    if title:
        r = cell.paragraphs[0].add_run(title)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x0D, 0x47, 0x2E)
    for line in content_lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run(line)
        r.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_table_with_header(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '0D472E')
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = 'F9FBF9' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def add_page_break(doc):
    doc.add_page_break()

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCDDCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)

# ─── Document Setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.0)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
style.paragraph_format.line_spacing = 1.15

# ─── COVER PAGE ────────────────────────────────────────────────────────────────

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
r = p.add_run('SILAKAP V1.0')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x6B, 0x9E, 0x7A)
r.font.bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
r = p.add_run('SIDATA')
r.font.size = Pt(42)
r.font.bold = True
r.font.color.rgb = RGBColor(0x0D, 0x47, 0x2E)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
r = p.add_run('Sistem Data Aparatur Sipil Negara')
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
r = p.add_run('Panduan Penggunaan Lengkap')
r.font.size = Pt(13)
r.font.italic = True
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Divider line
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('─' * 50)
r.font.color.rgb = RGBColor(0x0D, 0x47, 0x2E)

# Metadata box
table = doc.add_table(rows=4, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in table.rows:
    for cell in row.cells:
        set_cell_bg(cell, 'F0F8F4')

meta = [
    ('Versi Dokumen', 'v1.0 — Final'),
    ('Tanggal', datetime.date.today().strftime('%d %B %Y')),
    ('Berlaku Untuk', 'Seluruh Pengguna BKPSDM'),
    ('Status', 'APPROVED — Siap Digunakan'),
]
for i, (k, v) in enumerate(meta):
    table.rows[i].cells[0].paragraphs[0].add_run(k).font.bold = True
    table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    table.rows[i].cells[1].paragraphs[0].add_run(v).font.size = Pt(10)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
r = p.add_run('Badan Kepegawaian dan Pengembangan Sumber Daya Manusia')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_page_break(doc)

# ─── DAFTAR ISI ────────────────────────────────────────────────────────────────

add_heading(doc, 'DAFTAR ISI', level=1)
toc_items = [
    ('1.', 'Pendahuluan', '3'),
    ('2.', 'Gambaran Umum Modul SIDATA', '4'),
    ('3.', 'Hak Akses dan Peran Pengguna', '5'),
    ('4.', 'Dashboard SIDATA', '7'),
    ('5.', 'Direktori ASN', '9'),
    ('6.', 'Detail Profil ASN', '11'),
    ('7.', 'Pemutakhiran Data ASN — Pipeline 5 Langkah', '13'),
    ('8.', 'Import SIASN', '15'),
    ('9.', 'Import Referensi', '19'),
    ('10.', 'Mapping Referensi', '21'),
    ('11.', 'Riwayat Import', '24'),
    ('12.', 'Log Sinkronisasi', '25'),
    ('13.', 'Validasi Data', '26'),
    ('14.', 'Rekonsiliasi Data ASN', '28'),
    ('15.', 'Referensi Data Master', '30'),
    ('16.', 'Dokumen ASN', '32'),
    ('17.', 'Laporan SIDATA', '33'),
    ('18.', 'Konsep dan Terminologi', '35'),
    ('19.', 'Pertanyaan Umum (FAQ)', '37'),
    ('20.', 'Lampiran — Alur Kerja Lengkap', '39'),
]
toc_table = doc.add_table(rows=len(toc_items), cols=3)
for i, (num, title, pg) in enumerate(toc_items):
    toc_table.rows[i].cells[0].paragraphs[0].add_run(num).font.size = Pt(10)
    toc_table.rows[i].cells[1].paragraphs[0].add_run(title).font.size = Pt(10)
    r = toc_table.rows[i].cells[2].paragraphs[0].add_run(pg)
    r.font.size = Pt(10)
    toc_table.rows[i].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for ci, w in [(0, 1.0), (1, 12.0), (2, 1.5)]:
        toc_table.rows[i].cells[ci].width = Cm(w)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 1 — PENDAHULUAN
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 1 — PENDAHULUAN', level=1)

add_heading(doc, '1.1 Latar Belakang', level=2)
add_para(doc,
    'SIDATA (Sistem Data Aparatur Sipil Negara) adalah modul inti dalam platform SILAKAP V1.0 yang '
    'dikembangkan oleh Badan Kepegawaian dan Pengembangan Sumber Daya Manusia (BKPSDM). Modul ini '
    'berfungsi sebagai pangkalan data tunggal (single source of truth) bagi seluruh data kepegawaian '
    'ASN di lingkungan pemerintah daerah.')
add_para(doc,
    'Data yang dikelola melalui SIDATA menjadi dasar bagi seluruh modul lain dalam SILAKAP, termasuk '
    'SIAP (kinerja), SIPENSIUN (pensiun), SIFORMEN (formasi), dan modul analitik lainnya. Oleh karena '
    'itu, kualitas data di SIDATA secara langsung memengaruhi akurasi seluruh sistem.')

add_heading(doc, '1.2 Tujuan Dokumen', level=2)
add_para(doc, 'Dokumen ini disusun untuk:')
add_bullet(doc, 'Memandu seluruh pengguna SIDATA dalam mengoperasikan setiap fitur modul.')
add_bullet(doc, 'Menjelaskan alur kerja pemutakhiran data ASN dari sumber SIASN BKN hingga data master.')
add_bullet(doc, 'Menerangkan hak akses dan tanggung jawab setiap peran pengguna.')
add_bullet(doc, 'Menjadi referensi teknis operasional bagi operator, reviewer, dan administrator.')

add_heading(doc, '1.3 Ruang Lingkup', level=2)
add_para(doc, 'Dokumen ini mencakup seluruh fitur SIDATA versi 1.0, meliputi:')
add_bullet(doc, 'Dashboard dan monitoring kualitas data')
add_bullet(doc, 'Direktori dan profil ASN')
add_bullet(doc, 'Pipeline pemutakhiran data 5 langkah')
add_bullet(doc, 'Import data SIASN (PNS, PPPK, PPPK Paruh Waktu)')
add_bullet(doc, 'Manajemen referensi data master')
add_bullet(doc, 'Mapping, validasi, rekonsiliasi, dan commit data')
add_bullet(doc, 'Pengelolaan dokumen ASN')
add_bullet(doc, 'Laporan eksekutif SIDATA')

add_heading(doc, '1.4 Referensi', level=2)
add_bullet(doc, 'SOP DAT-002: Pemutakhiran Data ASN')
add_bullet(doc, 'SOP DAT-003: Rekonsiliasi Data ASN')
add_bullet(doc, 'SOP SIK-002: Verifikasi Data SIASN vs SIDATA')
add_bullet(doc, 'PermenPAN-RB No. 1 Tahun 2020 tentang Pedoman Penyusunan Kebutuhan Pegawai ASN')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 2 — GAMBARAN UMUM
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 2 — GAMBARAN UMUM MODUL SIDATA', level=1)

add_heading(doc, '2.1 Fungsi Utama', level=2)
add_para(doc,
    'SIDATA melayani enam fungsi utama dalam ekosistem SILAKAP:')

add_table_with_header(doc,
    ['No', 'Fungsi', 'Deskripsi'],
    [
        ['1', 'Master Data ASN', 'Menyimpan dan mengelola profil lengkap seluruh ASN aktif maupun nonaktif'],
        ['2', 'Import & Sinkronisasi', 'Mengimpor data dari SIASN BKN dan memvalidasinya sebelum masuk ke master'],
        ['3', 'Manajemen Referensi', 'Mengelola data referensi: jabatan, golongan, unit kerja, pangkat, dll.'],
        ['4', 'Validasi & Rekonsiliasi', 'Membandingkan data batch dengan master dan mengidentifikasi selisih'],
        ['5', 'Dokumen Kepegawaian', 'Mengunggah dan mengunduh dokumen pendukung per ASN'],
        ['6', 'Laporan & Analitik', 'Menghasilkan laporan kualitas data dan rekapitulasi eksekutif'],
    ],
    col_widths=[0.8, 4.0, 10.0]
)

add_heading(doc, '2.2 Arsitektur Data', level=2)
add_para(doc, 'Alur data dalam SIDATA mengikuti prinsip staged import:')
add_info_box(doc, 'Alur Data SIDATA', [
    '  SIASN BKN  →  [Upload Batch]  →  Staging Area  →  [Map + Validate]  →  Master ASN',
    '',
    '  Referensi BKN  →  [Upload Ref]  →  Staging Ref  →  [Commit]  →  Master Referensi',
    '',
    '  Master ASN  ←→  [Modul Lain: SIAP, SIPENSIUN, SIFORMEN, Analytics]',
], 'EAF4EE')

add_heading(doc, '2.3 Navigasi Menu SIDATA', level=2)
add_para(doc, 'Menu SIDATA dapat diakses melalui sidebar kiri aplikasi SILAKAP:')

add_table_with_header(doc,
    ['Menu', 'Submenu', 'Path', 'Fungsi Singkat'],
    [
        ['SIDATA ASN', 'Dashboard', '/sidata/dashboard', 'Monitoring kualitas & statistik'],
        ['', 'Profil ASN', '/sidata/asn', 'Direktori & pencarian ASN'],
        ['', 'Validasi', '/sidata/validasi', 'Pantau kualitas import'],
        ['', 'Pemutakhiran', '/sidata/pemutakhiran', 'Pusat kendali pipeline data'],
        ['Import Workspace', 'Import SIASN', '/sidata/import/siasn', 'Upload data ASN dari BKN'],
        ['', 'Import Referensi', '/sidata/import/referensi', 'Upload data referensi'],
        ['', 'Mapping Referensi', '/sidata/import/mapping-referensi', 'Petakan data ke referensi master'],
        ['', 'Riwayat Import', '/sidata/import/riwayat', 'Semua histori batch import'],
        ['', 'Log Sinkronisasi', '/sidata/import/log-sinkronisasi', 'Audit trail operasi'],
        ['', 'Rekonsiliasi', '/sidata/rekonsiliasi', 'Selisih batch vs master'],
        ['', 'Referensi Data', '/sidata/referensi', 'Kelola data referensi master'],
        ['', 'Dokumen ASN', '/sidata/dokumen', 'Upload/download dokumen ASN'],
        ['', 'Laporan', '/sidata/laporan', 'Laporan eksekutif SIDATA'],
    ],
    col_widths=[3.5, 3.5, 5.0, 4.5]
)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 3 — HAK AKSES
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 3 — HAK AKSES DAN PERAN PENGGUNA', level=1)

add_heading(doc, '3.1 Daftar Peran dan Izin', level=2)
add_para(doc,
    'SIDATA menggunakan sistem RBAC (Role-Based Access Control). Setiap pengguna memiliki '
    'satu atau lebih peran yang menentukan fitur yang dapat diakses.')

add_table_with_header(doc,
    ['Peran', 'Lihat', 'Input', 'Upload', 'Verifikasi', 'Review', 'Setujui', 'Commit', 'Admin'],
    [
        ['SUPER_ADMIN',         '✓', '✓', '✓', '✓', '✓', '✓', '✓', '✓'],
        ['ADMIN_BKPSDM',        '✓', '✓', '✓', '✓', '✓', '✓', '✓', '✓'],
        ['KEPALA_BADAN',        '✓', '—', '—', '—', '—', '—', '—', '—'],
        ['KABID',               '✓', '✓', '✓', '✓', '✓', '✓', '—', '—'],
        ['ANALIS_MADYA',        '✓', '✓', '✓', '✓', '✓', '—', '—', '—'],
        ['ANALIS_MUDA',         '✓', '✓', '✓', '✓', '—', '—', '—', '—'],
        ['ANALIS_PERTAMA',      '✓', '✓', '✓', '—', '—', '—', '—', '—'],
        ['PENELAAH',            '✓', '✓', '✓', '—', '—', '—', '—', '—'],
        ['PPPK',                '✓', '✓', '—', '—', '—', '—', '—', '—'],
        ['OPERATOR_IMPORT',     '✓', '—', '✓', '✓', '—', '—', '—', '—'],
        ['REVIEWER_MAPPING',    '✓', '—', '—', '—', '✓', '—', '—', '—'],
    ],
    col_widths=[4.0, 1.3, 1.3, 1.3, 1.8, 1.5, 1.5, 1.5, 1.5]
)

add_heading(doc, '3.2 Operasi Kritis dan Peran yang Diperlukan', level=2)
add_table_with_header(doc,
    ['Operasi', 'Peran Minimum yang Diizinkan'],
    [
        ['Upload file ASN/Referensi', 'SUPER_ADMIN, ADMIN_BKPSDM, OPERATOR_IMPORT'],
        ['Auto-map batch ASN', 'SUPER_ADMIN, ADMIN_BKPSDM, REVIEWER_MAPPING'],
        ['Commit batch ke master', 'SUPER_ADMIN, ADMIN_BKPSDM'],
        ['Commit referensi ke master', 'SUPER_ADMIN, ADMIN_BKPSDM'],
        ['Edit profil ASN individual', 'SUPER_ADMIN, ADMIN_BKPSDM, KABID, ANALIS_MADYA'],
        ['Upload dokumen ASN', 'Semua peran dengan akses SIDATA'],
        ['Lihat laporan eksekutif', 'KEPALA_BADAN dan semua admin/analisis'],
        ['Rekonsiliasi data', 'OPERATOR_IMPORT, REVIEWER_MAPPING, dan admin'],
    ],
    col_widths=[6.5, 10.0]
)

add_heading(doc, '3.3 Cara Login dan Akses SIDATA', level=2)
add_numbered(doc, 'Buka aplikasi SILAKAP di browser (Chrome/Firefox/Edge terbaru).')
add_numbered(doc, 'Masukkan username dan password yang telah diberikan administrator.')
add_numbered(doc, 'Setelah login, klik menu "SIDATA" di navigasi atas atau sidebar kiri.')
add_numbered(doc, 'Menu yang tampil disesuaikan secara otomatis dengan peran yang dimiliki.')

add_info_box(doc, 'CATATAN PENTING', [
    '• Jika menu SIDATA tidak muncul, hubungi administrator untuk verifikasi hak akses.',
    '• Sesi login akan berakhir otomatis setelah periode tidak aktif.',
    '• Setiap operasi dicatat di Log Sinkronisasi untuk keperluan audit.',
], 'FFF8E1')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 4 — DASHBOARD
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 4 — DASHBOARD SIDATA', level=1)
add_para(doc,
    'Dashboard adalah halaman utama SIDATA yang menampilkan ringkasan kondisi data ASN secara real-time. '
    'Halaman ini dapat diakses melalui menu SIDATA → Dashboard.')

add_heading(doc, '4.1 Kartu KPI Utama', level=2)
add_para(doc, 'Bagian atas dashboard menampilkan delapan kartu indikator utama:')
add_table_with_header(doc,
    ['Kartu', 'Deskripsi', 'Tindak Lanjut jika Bermasalah'],
    [
        ['Total ASN', 'Jumlah keseluruhan ASN dalam sistem', '—'],
        ['ASN Aktif', 'Jumlah ASN berstatus AKTIF', '—'],
        ['Quality Score (%)', 'Persentase data yang lengkap di seluruh field kritis', 'Lakukan pemutakhiran data'],
        ['BUP 12 Bulan', 'ASN yang akan pensiun dalam 12 bulan ke depan', 'Buat usulan pensiun via SIPENSIUN'],
        ['BUP Terlewat', 'ASN aktif yang sudah melewati batas usia pensiun', 'Segera proses pensiun'],
        ['Data Issue', 'Jumlah baris bermasalah di semua batch import', 'Buka Validasi Data'],
        ['Tanpa Unit Kerja', 'ASN tanpa unit kerja terdaftar', 'Lakukan remapping unit kerja'],
        ['Tanpa Jabatan', 'ASN tanpa jabatan terdaftar', 'Perbarui data jabatan via import'],
    ],
    col_widths=[3.5, 6.0, 7.0]
)

add_heading(doc, '4.2 Bagian Analisis Kualitas Data', level=2)
add_para(doc,
    'Di bawah KPI, terdapat panel "Kualitas Data" yang merinci ketidaklengkapan data per field:')
add_bullet(doc, 'Tanpa Unit Kerja — ASN yang belum memiliki unit kerja terdaftar')
add_bullet(doc, 'Tanpa Jabatan — ASN yang belum memiliki jabatan terdaftar')
add_bullet(doc, 'Tanpa Golongan — ASN yang belum memiliki golongan/pangkat terdaftar')
add_bullet(doc, 'Tanpa NIK — ASN yang belum memiliki Nomor Induk Kependudukan')
add_bullet(doc, 'Tanpa Tanggal Lahir — ASN yang data tanggal lahirnya kosong')
add_bullet(doc, 'Tanpa TMT Pensiun — ASN yang belum memiliki tanggal mulai pensiun')
add_bullet(doc, 'Tanpa Profil SIASN — ASN yang belum tersinkron dengan sistem BKN')

add_heading(doc, '4.3 Monitor Batch Import Terbaru', level=2)
add_para(doc,
    'Tabel di bagian bawah dashboard menampilkan 8 batch import terbaru beserta statusnya. '
    'Setiap baris menampilkan:')
add_bullet(doc, 'ID Batch — kode unik identifikasi batch')
add_bullet(doc, 'Jenis — ASN atau Referensi')
add_bullet(doc, 'Nama File — file sumber yang diunggah')
add_bullet(doc, 'Status — UPLOADED / MAPPED / COMMITTED / FAILED')
add_bullet(doc, 'Jumlah baris: Total, Valid, Invalid, Warning')
add_bullet(doc, 'Tanggal upload')
add_bullet(doc, 'Tombol "Buka" — navigasi ke detail batch')

add_heading(doc, '4.4 Aksi Cepat (Quick Actions)', level=2)
add_para(doc, 'Panel aksi cepat di dashboard memudahkan navigasi ke halaman yang sering digunakan:')
add_bullet(doc, 'Buka Profil ASN → ke halaman Direktori ASN')
add_bullet(doc, 'Import Data SIASN → ke halaman Import SIASN')
add_bullet(doc, 'Import Referensi → ke halaman Import Referensi')
add_bullet(doc, 'Review Mapping → ke halaman Mapping Referensi')
add_bullet(doc, 'Riwayat Import → ke halaman Riwayat Import')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 5 — DIREKTORI ASN
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 5 — DIREKTORI ASN', level=1)
add_para(doc,
    'Halaman Direktori ASN (menu: SIDATA → Profil ASN) menampilkan seluruh database ASN '
    'dalam bentuk tabel yang dapat dicari dan difilter.')

add_heading(doc, '5.1 Pencarian dan Filter', level=2)
add_para(doc, 'Tersedia empat opsi filter yang dapat dikombinasikan:')
add_table_with_header(doc,
    ['Filter', 'Cara Penggunaan', 'Contoh'],
    [
        ['Pencarian Teks', 'Ketik NIP, NIK, nama, atau nama jabatan. Hasil muncul otomatis setelah 650ms.', '"Budi Santoso" atau "197801012005011001"'],
        ['Status ASN', 'Pilih dari dropdown: Aktif, Pensiun, CLTN, Berhenti, Meninggal', 'Pilih "Aktif" untuk tampilkan ASN aktif saja'],
        ['Jenis ASN', 'Pilih: PNS, PPPK, atau PPPK Paruh Waktu', 'Pilih "PPPK" untuk data kontrak'],
        ['Unit Kerja', 'Klik ikon pohon organisasi, pilih unit dari hierarki yang dapat diperluas', 'Pilih "Bidang Pengadaan ASN"'],
    ],
    col_widths=[3.5, 7.5, 5.5]
)

add_heading(doc, '5.2 Kolom Tabel', level=2)
add_table_with_header(doc,
    ['Kolom', 'Isi', 'Keterangan'],
    [
        ['NIP', 'Nomor Induk Pegawai', 'Format 18 digit'],
        ['Nama', 'Nama lengkap + jabatan + usia/pendidikan', 'Klik tombol Detail untuk profil penuh'],
        ['Jenis', 'Badge PNS / PPPK / PPPK PW', 'Warna berbeda per jenis'],
        ['Unit Kerja', 'Nama unit kerja', '—'],
        ['Golongan', 'Nama golongan + TMT + masa kerja', '—'],
        ['Status', 'Badge status kepegawaian', 'Hijau = Aktif, Abu = Nonaktif'],
        ['Aksi', 'Tombol Detail dan Usulan', '"Usulan" hanya muncul untuk status AKTIF'],
    ],
    col_widths=[2.5, 5.5, 8.5]
)

add_heading(doc, '5.3 Ekspor Data ke Excel', level=2)
add_para(doc, 'Untuk mengunduh data ASN dalam format Excel:')
add_numbered(doc, 'Atur filter sesuai kebutuhan (opsional).')
add_numbered(doc, 'Klik tombol "Export Excel" di kanan atas tabel.')
add_numbered(doc, 'File .xlsx akan diunduh otomatis ke komputer Anda.')
add_numbered(doc, 'File berisi seluruh data ASN sesuai filter yang aktif (bukan hanya halaman yang ditampilkan).')

add_info_box(doc, 'INFO', [
    '• Ekspor bisa memakan waktu beberapa detik jika jumlah ASN sangat banyak.',
    '• Kolom yang diekspor mencakup: NIP, NIK, Nama, Jenis ASN, Status, Unit Kerja, Jabatan, Golongan, TMT Pensiun.',
], 'E3F2FD')

add_heading(doc, '5.4 Paginasi', level=2)
add_para(doc, 'Data ditampilkan 20 baris per halaman. Gunakan tombol "Sebelumnya" dan "Berikutnya" '
    'di bawah tabel untuk berpindah halaman. Indikator halaman menampilkan posisi baris saat ini.')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 6 — DETAIL PROFIL ASN
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 6 — DETAIL PROFIL ASN', level=1)
add_para(doc,
    'Halaman Detail ASN dapat diakses dengan mengklik tombol "Detail" pada baris ASN di Direktori. '
    'Halaman ini menampilkan profil lengkap satu ASN dalam tiga tab.')

add_heading(doc, '6.1 Tab Profil', level=2)
add_para(doc, 'Tab default yang menampilkan data kepegawaian dan identitas:')
add_bullet(doc, 'Data Kepegawaian: NIP, Jabatan, Golongan, Jenis ASN, Status Kepegawaian, Unit Kerja, TMT Pensiun')
add_bullet(doc, 'Identitas & Kontak: Tanggal Lahir, Email, Nomor Telepon')
add_bullet(doc, 'Tombol "Buat Usulan Pensiun" — hanya muncul jika status ASN = AKTIF')

add_heading(doc, '6.2 Tab Riwayat', level=2)
add_para(doc, 'Menampilkan histori perubahan data ASN yang dimuat secara lazy (on demand):')
add_bullet(doc, 'Riwayat Jabatan & Unit Kerja: daftar perubahan jabatan dan unit dengan tanggal, sumber batch, dan ID sinkronisasi')
add_bullet(doc, 'Riwayat Golongan: daftar perubahan golongan/pangkat dengan tanggal dan sumber')
add_bullet(doc, 'Jika belum ada riwayat, ditampilkan pesan "Belum ada riwayat"')

add_heading(doc, '6.3 Tab Edit', level=2)
add_para(doc, 'Memungkinkan pengeditan langsung field tertentu (memerlukan izin write):')
add_table_with_header(doc,
    ['Field yang Dapat Diedit', 'Keterangan'],
    [
        ['Nama', 'Nama lengkap ASN'],
        ['NIK', 'Nomor Induk Kependudukan 16 digit'],
        ['Jabatan', 'Pilih dari referensi jabatan yang tersedia'],
        ['Golongan', 'Pilih dari referensi golongan yang tersedia'],
        ['Status ASN', 'Ubah status: AKTIF / PENSIUN / CLTN / BERHENTI / MENINGGAL'],
        ['TMT Pensiun', 'Tanggal mulai pensiun (format YYYY-MM-DD)'],
    ],
    col_widths=[5.5, 11.0]
)
add_info_box(doc, 'PERHATIAN', [
    '• Perubahan via Tab Edit tidak melalui pipeline dan tidak memiliki staging.',
    '• Gunakan edit langsung hanya untuk koreksi data minor yang tidak memerlukan audit batch.',
    '• Untuk pemutakhiran massal, gunakan Pipeline Pemutakhiran Data (Bab 7).',
], 'FFF3E0')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 7 — PIPELINE PEMUTAKHIRAN
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 7 — PEMUTAKHIRAN DATA ASN — PIPELINE 5 LANGKAH', level=1)
add_para(doc,
    'Pemutakhiran data ASN adalah proses resmi untuk menyinkronkan data SIASN BKN dengan '
    'master SIDATA. Proses ini dirancang dengan 5 langkah berurutan untuk memastikan kualitas '
    'dan auditabilitas data.')

add_heading(doc, '7.1 Gambaran Pipeline', level=2)
add_info_box(doc, 'Alur Pipeline Pemutakhiran Data', [
    '',
    '  [1] Upload SIASN  →  [2] Mapping Referensi  →  [3] Validasi Data',
    '                                                          ↓',
    '            [5] Audit Log  ←  [4] Commit Data  ←────────┘',
    '',
    '  Setiap langkah harus diselesaikan sebelum melanjutkan ke langkah berikutnya.',
], 'EAF4EE')

add_heading(doc, '7.2 Langkah 1 — Upload SIASN', level=2)
add_para(doc, 'Tujuan: Mengunggah file Excel data ASN dari sistem SIASN BKN.')
add_numbered(doc, 'Buka menu SIDATA → Pemutakhiran, atau langsung ke Import SIASN.')
add_numbered(doc, 'Pilih Tipe Pegawai: PNS, PPPK, atau PPPK Paruh Waktu.')
add_numbered(doc, 'Klik tombol "Pilih File" dan pilih file Excel (.xlsx) dari SIASN.')
add_numbered(doc, 'Klik "Upload". Sistem akan memproses file dan membuat batch baru.')
add_numbered(doc, 'Tunggu hingga status batch berubah menjadi "VALIDATED".')
add_para(doc, 'Peran yang diizinkan: SUPER_ADMIN, ADMIN_BKPSDM, OPERATOR_IMPORT')

add_heading(doc, '7.3 Langkah 2 — Mapping Referensi', level=2)
add_para(doc, 'Tujuan: Memetakan nilai teks dalam file SIASN ke referensi master di SIDATA.')
add_numbered(doc, 'Buka menu Import Workspace → Mapping Referensi.')
add_numbered(doc, 'Pilih batch yang baru diunggah dari dropdown.')
add_numbered(doc, 'Klik tombol "Auto-Map" untuk menjalankan pemetaan otomatis.')
add_numbered(doc, 'Tinjau baris yang berstatus "Perlu Review" — baris ini memerlukan konfirmasi manual.')
add_numbered(doc, 'Untuk baris Unmapped: cari padanan manual dari daftar referensi yang tersedia.')
add_numbered(doc, 'Jika ada unit kerja yang tidak dikenali, gunakan fitur "Resolve Unit Mapping".')
add_numbered(doc, 'Setelah semua baris terpetakan, klik "Remap" untuk memperbarui hasil pemetaan.')
add_para(doc, 'Peran yang diizinkan: SUPER_ADMIN, ADMIN_BKPSDM, REVIEWER_MAPPING')

add_heading(doc, '7.4 Langkah 3 — Validasi Data', level=2)
add_para(doc, 'Tujuan: Memeriksa kualitas data sebelum dicommit ke master.')
add_numbered(doc, 'Buka menu SIDATA → Validasi Data.')
add_numbered(doc, 'Periksa ringkasan kualitas (skor persentase dan jumlah baris bermasalah).')
add_numbered(doc, 'Tinjau tab "Invalid" untuk baris yang gagal validasi struktural.')
add_numbered(doc, 'Tinjau tab "Perlu Review" untuk baris yang memerlukan persetujuan manual.')
add_numbered(doc, 'Unduh laporan issue jika diperlukan untuk analisis lebih lanjut.')
add_numbered(doc, 'Tentukan apakah batch layak dilanjutkan ke commit atau perlu diunggah ulang.')

add_heading(doc, '7.5 Langkah 4 — Commit Data', level=2)
add_para(doc, 'Tujuan: Menerapkan data batch ke master ASN secara permanen.')
add_info_box(doc, 'PERINGATAN — Operasi Tidak Dapat Dibatalkan', [
    '• Commit adalah operasi FINAL. Data yang sudah dicommit tidak dapat dirollback.',
    '• Pastikan validasi sudah selesai dan semua issue yang kritis sudah diselesaikan.',
    '• Hanya SUPER_ADMIN dan ADMIN_BKPSDM yang dapat melakukan commit.',
], 'FFEBEE')
add_numbered(doc, 'Buka detail batch di Riwayat Import atau melalui link di halaman Mapping.')
add_numbered(doc, 'Klik tombol "Commit Batch".')
add_numbered(doc, 'Konfirmasi dialog peringatan.')
add_numbered(doc, 'Tunggu proses selesai — status batch akan berubah menjadi "COMMITTED".')
add_numbered(doc, 'Data master ASN diperbarui sesuai isi batch.')

add_heading(doc, '7.6 Langkah 5 — Audit Log', level=2)
add_para(doc, 'Tujuan: Memverifikasi bahwa semua operasi tercatat dengan benar.')
add_numbered(doc, 'Buka menu Import Workspace → Log Sinkronisasi.')
add_numbered(doc, 'Filter berdasarkan batch atau jenis aksi.')
add_numbered(doc, 'Verifikasi setiap langkah operasi tercatat dengan user, timestamp, dan hasil.')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 8 — IMPORT SIASN
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 8 — IMPORT SIASN', level=1)
add_para(doc,
    'Halaman Import SIASN (menu: Import Workspace → Import SIASN) adalah antarmuka utama '
    'untuk mengunggah file data ASN dari sistem SIASN BKN.')

add_heading(doc, '8.1 Jenis File yang Didukung', level=2)
add_table_with_header(doc,
    ['Tipe Pegawai', 'Kode Import', 'Keterangan', 'Warna Identifikasi'],
    [
        ['PNS', 'SIASN_ASN_PNS', 'Pegawai Negeri Sipil penuh waktu', 'Biru'],
        ['PPPK', 'SIASN_ASN_PPPK', 'Pegawai Pemerintah Perjanjian Kerja', 'Ungu'],
        ['PPPK Paruh Waktu', 'SIASN_ASN_PPPK_PARUH_WAKTU', 'PPPK dengan jam kerja paruh waktu', 'Oranye'],
    ],
    col_widths=[3.5, 5.0, 6.0, 3.0]
)

add_heading(doc, '8.2 Langkah Upload', level=2)
add_numbered(doc, 'Pilih Tipe Pegawai dengan mengklik salah satu dari tiga tombol (PNS / PPPK / PPPK Paruh Waktu).')
add_numbered(doc, 'Klik area upload atau tombol "Pilih File". File harus berformat .xlsx (Excel).')
add_numbered(doc, 'Klik tombol "Upload". Sistem menerapkan throttle 5 permintaan per menit.')
add_numbered(doc, 'Proses upload akan menampilkan progress. Hasil upload berupa:')
add_bullet(doc, 'Total baris yang diproses', level=1)
add_bullet(doc, 'Baris valid, invalid, warning, dan perlu review', level=1)
add_bullet(doc, 'ID Batch yang dihasilkan', level=1)
add_numbered(doc, 'Setelah upload selesai, batch muncul di daftar batch di bawah halaman.')

add_heading(doc, '8.3 Status Batch Import', level=2)
add_table_with_header(doc,
    ['Status', 'Arti', 'Tindakan Selanjutnya'],
    [
        ['UPLOADED', 'File diterima, validasi awal selesai', 'Lanjut ke Mapping Referensi'],
        ['PROCESSING', 'Sistem sedang memproses file', 'Tunggu hingga selesai'],
        ['VALIDATED', 'Validasi format dan isi selesai', 'Tinjau issue lalu lanjut mapping'],
        ['MAPPED', 'Pemetaan referensi berhasil dijalankan', 'Tinjau unmapped, lalu commit'],
        ['COMMITTED', 'Data sudah diterapkan ke master', 'Tidak ada tindakan lebih lanjut'],
        ['FAILED', 'Proses gagal (format tidak valid, dll.)', 'Periksa error, upload ulang'],
        ['CANCELLED', 'Batch dibatalkan oleh pengguna', 'Upload file baru jika diperlukan'],
    ],
    col_widths=[3.0, 6.0, 7.5]
)

add_heading(doc, '8.4 Meninjau Issue dalam Batch', level=2)
add_para(doc, 'Setelah upload, tab issue di bagian bawah halaman menampilkan baris bermasalah:')
add_bullet(doc, 'Semua Issue — gabungan semua kategori masalah')
add_bullet(doc, 'Perlu Review — baris yang perlu konfirmasi manual sebelum commit')
add_bullet(doc, 'Invalid — baris yang gagal validasi dan tidak dapat dicommit')
add_para(doc, 'Untuk setiap issue, ditampilkan: nomor baris, NIP/NIK, nama, kode error, dan deskripsi masalah.')

add_heading(doc, '8.5 Contoh Error Umum', level=2)
add_table_with_header(doc,
    ['Kode Error', 'Penyebab', 'Cara Mengatasi'],
    [
        ['DUPLICATE_NIP', 'NIP sudah ada di master dengan data berbeda', 'Verifikasi data sumber di SIASN'],
        ['MISSING_UNIT', 'Unit kerja tidak ditemukan di referensi', 'Tambah unit kerja di Referensi Data atau resolve mapping'],
        ['INVALID_GOLONGAN', 'Kode golongan tidak dikenali sistem', 'Import referensi golongan terbaru'],
        ['MISSING_JABATAN', 'Jabatan tidak ditemukan di referensi', 'Import referensi jabatan atau peta secara manual'],
        ['INVALID_DATE', 'Format tanggal tidak sesuai (bukan YYYY-MM-DD)', 'Perbaiki format tanggal di file sumber'],
    ],
    col_widths=[3.5, 5.5, 7.5]
)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 9 — IMPORT REFERENSI
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 9 — IMPORT REFERENSI', level=1)
add_para(doc,
    'Halaman Import Referensi (menu: Import Workspace → Import Referensi) digunakan untuk '
    'memperbarui data referensi master yang digunakan oleh proses mapping ASN.')

add_heading(doc, '9.1 Jenis Referensi yang Didukung', level=2)
add_table_with_header(doc,
    ['Jenis Referensi', 'Deskripsi', 'Endpoint'],
    [
        ['Jabatan Struktural', 'Jabatan Eselon I-IV', 'POST /reference-jabatan/upload (jenisJabatan=STRUKTURAL)'],
        ['Jabatan Fungsional', 'Jabatan fungsional umum & tertentu', 'POST /reference-jabatan/upload (jenisJabatan=FUNGSIONAL)'],
        ['Jabatan Pelaksana', 'Jabatan pelaksana/staf', 'POST /reference-jabatan/upload (jenisJabatan=PELAKSANA)'],
        ['Golongan', 'Golongan ruang (I/a sampai IV/e)', 'POST /reference/upload (referenceType=GOLONGAN)'],
        ['Pangkat', 'Nama pangkat sesuai golongan', 'POST /reference/upload (referenceType=PANGKAT)'],
        ['Pendidikan', 'Tingkat pendidikan (SD hingga S3)', 'POST /reference/upload (referenceType=PENDIDIKAN)'],
        ['Agama', 'Daftar agama yang diakui', 'POST /reference/upload (referenceType=AGAMA)'],
        ['Jenis Kelamin', 'Laki-laki / Perempuan', 'POST /reference/upload (referenceType=JENIS_KELAMIN)'],
        ['Status Kawin', 'Belum Kawin / Kawin / Cerai', 'POST /reference/upload (referenceType=STATUS_KAWIN)'],
        ['Kedudukan Hukum', 'Status hukum ASN', 'POST /reference/upload (referenceType=KEDUDUKAN_HUKUM)'],
        ['Jenis ASN', 'PNS / PPPK / dst.', 'POST /reference/upload (referenceType=JENIS_ASN)'],
    ],
    col_widths=[3.5, 5.0, 8.0]
)

add_heading(doc, '9.2 Langkah Import Referensi', level=2)
add_numbered(doc, 'Buka halaman Import Referensi.')
add_numbered(doc, 'Pilih jenis referensi yang akan diperbarui.')
add_numbered(doc, 'Siapkan file Excel sesuai format template referensi (kolom: kode, nama).')
add_numbered(doc, 'Klik "Pilih File" dan pilih file .xlsx yang sesuai.')
add_numbered(doc, 'Klik "Upload". Sistem akan membuat batch referensi baru.')
add_numbered(doc, 'Setelah batch terbuat, klik "Commit" untuk menerapkan ke master referensi.')

add_info_box(doc, 'FORMAT FILE REFERENSI', [
    '• Kolom minimal: kode (varchar 100) dan nama (varchar 300)',
    '• Untuk Jabatan: tambah kolom jenisJabatanId',
    '• Baris header wajib ada di baris pertama',
    '• Format file: .xlsx (Microsoft Excel 2007+)',
], 'E8F5E9')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 10 — MAPPING REFERENSI
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 10 — MAPPING REFERENSI', level=1)
add_para(doc,
    'Mapping Referensi adalah proses memetakan nilai teks dalam file SIASN (misalnya nama unit kerja '
    '"Bidang I") ke entri referensi master yang terdaftar di sistem.')

add_heading(doc, '10.1 Konsep Mapping', level=2)
add_para(doc,
    'Ketika file SIASN diunggah, baris-baris di dalamnya berisi nilai teks seperti nama jabatan, '
    'nama golongan, dan nama unit kerja. Sistem perlu "mencocokkan" nilai-nilai ini dengan ID referensi '
    'internal di database. Proses inilah yang disebut mapping.')
add_info_box(doc, 'Contoh Mapping', [
    '  File SIASN berisi: jabatan = "Analis Kepegawaian Madya"',
    '  Sistem mencari di master referensi jabatan...',
    '  → Ditemukan: RefJabatan.id = "abc-123", nama = "Analis Kepegawaian"',
    '  → Status: MAPPED (siap dicommit)',
    '',
    '  File SIASN berisi: unitKerja = "Bidang Pengadaan"',
    '  Sistem mencari... tidak ditemukan padanan.',
    '  → Status: UNMAPPED (perlu penyelesaian manual)',
], 'EAF4EE')

add_heading(doc, '10.2 Status Mapping per Baris', level=2)
add_table_with_header(doc,
    ['Status', 'Arti', 'Tindakan'],
    [
        ['MAPPED', 'Padanan ditemukan otomatis', 'Tidak perlu tindakan'],
        ['NEEDS_REVIEW', 'Padanan ditemukan tapi confidence rendah', 'Konfirmasi atau tolak padanan yang disarankan'],
        ['UNMAPPED', 'Tidak ada padanan yang ditemukan', 'Cari dan pilih padanan manual'],
        ['INVALID', 'Data tidak valid secara struktural', 'Perbaiki file sumber dan upload ulang'],
    ],
    col_widths=[3.0, 6.5, 7.0]
)

add_heading(doc, '10.3 Langkah-langkah Mapping', level=2)
add_numbered(doc, 'Buka menu Import Workspace → Mapping Referensi.')
add_numbered(doc, 'Pilih batch dari dropdown "Pilih Batch". Hanya batch dengan status VALIDATED atau MAPPED yang ditampilkan.')
add_numbered(doc, 'Klik "Auto-Map" — sistem akan menjalankan pemetaan otomatis berbasis kemiripan nama.')
add_numbered(doc, 'Tinjau tab "Perlu Review": klik setiap baris dan konfirmasi atau tolak padanan yang disarankan.')
add_numbered(doc, 'Tinjau tab "Unmapped": untuk setiap baris, gunakan kotak pencarian untuk menemukan referensi yang tepat.')
add_numbered(doc, 'Jika ada unit kerja yang tidak dikenali:')
add_bullet(doc, 'Klik "Resolve Unit Mapping" pada baris yang bersangkutan', level=1)
add_bullet(doc, 'Cari unit kerja yang sesuai dari daftar', level=1)
add_bullet(doc, 'Tambahkan catatan jika diperlukan', level=1)
add_numbered(doc, 'Setelah semua baris diselesaikan, klik "Remap" untuk memperbarui status batch.')
add_numbered(doc, 'Verifikasi di ringkasan: semua baris seharusnya berstatus MAPPED atau COMMITTED.')

add_heading(doc, '10.4 Mengunduh Laporan Issue', level=2)
add_para(doc, 'Klik tombol "Download Issues" untuk mengunduh laporan baris bermasalah ke format Excel. '
    'File ini berguna untuk koordinasi dengan SIASN BKN jika ada data yang perlu diperbaiki di sumber.')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 11 — RIWAYAT IMPORT
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 11 — RIWAYAT IMPORT', level=1)
add_para(doc,
    'Halaman Riwayat Import (menu: Import Workspace → Riwayat Import) menampilkan seluruh histori '
    'batch import yang pernah dibuat, baik untuk ASN maupun referensi.')

add_heading(doc, '11.1 Fitur Pencarian dan Filter', level=2)
add_bullet(doc, 'Pencarian teks — cari berdasarkan ID batch, nama file, jenis import, jenis referensi')
add_bullet(doc, 'Filter Status — tampilkan batch dengan status tertentu')
add_bullet(doc, 'Filter Jenis — pilih antara batch ASN atau Referensi')

add_heading(doc, '11.2 Kolom Tabel Riwayat', level=2)
add_table_with_header(doc,
    ['Kolom', 'Isi'],
    [
        ['Batch ID', 'Kode unik batch (klik untuk detail)'],
        ['Jenis', 'ASN (PNS/PPPK/PPPK PW) atau Referensi (jenis ref)'],
        ['File', 'Nama file yang diunggah'],
        ['Status', 'Badge status terkini'],
        ['Total / Valid / Invalid', 'Jumlah baris per kategori'],
        ['Mapped / Review / Unmapped', 'Hasil proses mapping'],
        ['Committed', 'Jumlah baris yang sudah dicommit ke master'],
        ['Tanggal', 'Waktu batch dibuat'],
        ['Aksi', 'Tombol "Buka Workspace" untuk masuk ke detail batch'],
    ],
    col_widths=[2.5, 14.0]
)

add_heading(doc, '11.3 Membuka Detail Batch', level=2)
add_para(doc, 'Klik tombol "Buka" atau "Buka Workspace" pada baris batch untuk membuka tampilan '
    'detail batch yang menampilkan semua baris dan status mapping-nya.')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 12 — LOG SINKRONISASI
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 12 — LOG SINKRONISASI (AUDIT TRAIL)', level=1)
add_para(doc,
    'Log Sinkronisasi mencatat setiap operasi yang dilakukan di SIDATA — termasuk upload, mapping, '
    'commit, dan pembatalan batch. Log ini tidak dapat diubah dan bersifat permanen.')

add_heading(doc, '12.1 Jenis Aksi yang Dicatat', level=2)
add_table_with_header(doc,
    ['Kode Aksi', 'Deskripsi'],
    [
        ['UPLOAD_ASN', 'Pengunggahan file data ASN'],
        ['MAP_ASN', 'Eksekusi auto-map pada batch ASN'],
        ['REMAP_ASN', 'Eksekusi ulang mapping batch ASN'],
        ['COMMIT_ASN', 'Commit batch ASN ke master'],
        ['UPLOAD_REFERENCE', 'Pengunggahan file referensi'],
        ['COMMIT_REFERENCE', 'Commit referensi ke master'],
        ['VIEW_ISSUES', 'Pengguna membuka halaman review issue'],
        ['CANCEL_BATCH', 'Pembatalan batch oleh pengguna'],
    ],
    col_widths=[4.5, 12.0]
)

add_heading(doc, '12.2 Cara Membaca Log', level=2)
add_para(doc, 'Setiap entri log menampilkan:')
add_bullet(doc, 'Identitas pengguna (nama + peran) yang melakukan aksi')
add_bullet(doc, 'Jenis aksi dengan ikon dan warna status')
add_bullet(doc, 'Jenis batch (ASN/Referensi) dan ID batch terkait')
add_bullet(doc, 'Ringkasan hasil (baris berhasil, gagal, dll.)')
add_bullet(doc, 'Timestamp (tanggal dan waktu tepat)')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 13 — VALIDASI DATA
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 13 — VALIDASI DATA', level=1)
add_para(doc,
    'Halaman Validasi (menu: SIDATA → Validasi) adalah pusat pemantauan kualitas data import. '
    'Halaman ini mengidentifikasi batch yang memiliki masalah dan membutuhkan tindak lanjut.')

add_heading(doc, '13.1 Skor Kualitas', level=2)
add_para(doc, 'Bagian atas menampilkan skor kualitas data secara keseluruhan:')
add_table_with_header(doc,
    ['Rentang Skor', 'Status', 'Warna', 'Interpretasi'],
    [
        ['85% – 100%', 'Baik', 'Hijau', 'Data dalam kondisi ideal'],
        ['60% – 84%', 'Perlu Perhatian', 'Kuning', 'Ada beberapa field kritis yang kosong'],
        ['0% – 59%', 'Kritis', 'Merah', 'Segera lakukan pemutakhiran data'],
    ],
    col_widths=[3.5, 3.5, 2.5, 7.0]
)

add_heading(doc, '13.2 Batch Bermasalah', level=2)
add_para(doc,
    'Tabel "Batch Bermasalah" menampilkan semua batch yang memiliki satu atau lebih dari kondisi berikut:')
add_bullet(doc, 'Baris Invalid > 0')
add_bullet(doc, 'Baris Warning > 0')
add_bullet(doc, 'Baris Perlu Review > 0')
add_bullet(doc, 'Baris Unmapped > 0')
add_bullet(doc, 'Status = FAILED')
add_para(doc, 'Klik "Tindak Lanjut" pada baris batch untuk langsung masuk ke workspace mapping batch tersebut.')

add_heading(doc, '13.3 Sampling ASN', level=2)
add_para(doc,
    'Bagian "Sampling ASN" menampilkan 10 data ASN dari halaman pertama master. '
    'Sampling ini berguna untuk verifikasi cepat bahwa data commit terakhir sudah masuk dengan benar.')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 14 — REKONSILIASI DATA ASN
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 14 — REKONSILIASI DATA ASN', level=1)
add_para(doc,
    'Rekonsiliasi adalah proses membandingkan data dalam satu batch SIASN dengan kondisi master SIDATA '
    'saat ini, untuk mengidentifikasi perbedaan sebelum data dicommit.')

add_heading(doc, '14.1 Jenis Selisih (Reconciliation Type)', level=2)
add_table_with_header(doc,
    ['Tipe', 'Arti', 'Warna', 'Tindakan'],
    [
        ['BEDA_DATA', 'ASN ada di batch dan master, tapi ada field yang berbeda', 'Kuning', 'Tinjau selisih, konfirmasi apakah data batch lebih baru'],
        ['HANYA_DI_BATCH', 'ASN ada di batch tapi tidak ada di master (data baru)', 'Oranye', 'Verifikasi data baru, commit jika valid'],
        ['HANYA_DI_MASTER', 'ASN ada di master tapi tidak ada di batch (mungkin pensiun/keluar)', 'Merah', 'Verifikasi apakah ASN memang sudah tidak aktif'],
        ['SAMA', 'Data batch dan master identik', 'Hijau', 'Tidak perlu tindakan'],
    ],
    col_widths=[3.5, 5.5, 2.0, 5.5]
)

add_heading(doc, '14.2 Cara Menggunakan Rekonsiliasi', level=2)
add_numbered(doc, 'Buka menu Import Workspace → Rekonsiliasi.')
add_numbered(doc, 'Pilih batch dari dropdown "Pilih Batch".')
add_numbered(doc, 'Pilih filter tipe rekonsiliasi:')
add_bullet(doc, '"Semua Perlu Perhatian" — tampilkan semua selisih (default)', level=1)
add_bullet(doc, '"Beda Data" — hanya ASN dengan perbedaan field', level=1)
add_bullet(doc, '"Hanya di Batch" — ASN baru yang belum ada di master', level=1)
add_bullet(doc, '"Hanya di Master" — ASN yang ada di master tapi absen dari batch', level=1)
add_numbered(doc, 'Gunakan kotak pencarian untuk menyaring NIP, nama, atau unit kerja.')
add_numbered(doc, 'Klik baris untuk melihat detail perbedaan per field (nilai master vs nilai batch).')
add_numbered(doc, 'Setelah rekonsiliasi selesai, lanjutkan ke proses commit jika semua selisih sudah ditangani.')

add_heading(doc, '14.3 Ringkasan Statistik Rekonsiliasi', level=2)
add_para(doc, 'Bagian atas halaman rekonsiliasi menampilkan delapan kartu statistik:')
add_bullet(doc, 'Perlu Perhatian — total baris yang membutuhkan tinjauan (bukan SAMA)')
add_bullet(doc, 'Beda Data — jumlah ASN dengan perbedaan field')
add_bullet(doc, 'Hanya di Batch — jumlah data baru di batch')
add_bullet(doc, 'Hanya di Master — jumlah ASN yang absen dari batch')
add_bullet(doc, 'Batch Rows — total baris di batch')
add_bullet(doc, 'Master Rows — total ASN di master yang relevan')
add_bullet(doc, 'Sama — jumlah baris yang identik')
add_bullet(doc, 'Ditampilkan — jumlah baris yang saat ini terlihat sesuai filter')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 15 — REFERENSI DATA MASTER
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 15 — REFERENSI DATA MASTER', level=1)
add_para(doc,
    'Halaman Referensi Data (menu: SIDATA → Referensi Data) menampilkan dan memungkinkan '
    'pengelolaan 12 kategori data referensi yang digunakan oleh proses mapping ASN.')

add_heading(doc, '15.1 Kategori Referensi', level=2)
add_table_with_header(doc,
    ['No', 'Kategori', 'Keterangan', 'Dapat Tambah Manual'],
    [
        ['1', 'Jenis Jabatan', 'Struktural, Fungsional, Pelaksana', 'Tidak'],
        ['2', 'Jabatan', 'Daftar lengkap jabatan ASN', 'Ya (dengan jenisJabatanId)'],
        ['3', 'Unit Organisasi', 'Hierarki unit kerja OPD', 'Melalui Import'],
        ['4', 'Golongan', 'I/a sampai IV/e', 'Via Import Referensi'],
        ['5', 'Pangkat', 'Nama pangkat per golongan', 'Via Import Referensi'],
        ['6', 'Pendidikan', 'SD, SMP, SMA, D3, S1, S2, S3', 'Via Import Referensi'],
        ['7', 'Agama', 'Islam, Kristen, Katolik, dll.', 'Via Import Referensi'],
        ['8', 'Jenis Kelamin', 'Laki-laki, Perempuan', 'Via Import Referensi'],
        ['9', 'Status Kawin', 'Belum Kawin, Kawin, Cerai', 'Via Import Referensi'],
        ['10', 'Kedudukan Hukum', 'Status hukum pegawai', 'Via Import Referensi'],
        ['11', 'Jenis ASN', 'PNS, PPPK, dll.', 'Via Import Referensi'],
        ['12', 'Status ASN', 'Aktif, Pensiun, CLTN, dll.', 'Tidak (sistem)'],
    ],
    col_widths=[0.8, 3.5, 6.0, 4.0]
)

add_heading(doc, '15.2 Menambah Referensi Manual', level=2)
add_para(doc, 'Khusus untuk kategori Jabatan, penambahan manual dapat dilakukan:')
add_numbered(doc, 'Pilih tab "Jabatan" di halaman Referensi Data.')
add_numbered(doc, 'Isi kolom "Kode" (kode unik jabatan).')
add_numbered(doc, 'Isi kolom "Nama" (nama jabatan).')
add_numbered(doc, 'Pilih Jenis Jabatan (jenisJabatanId).')
add_numbered(doc, 'Klik "Tambah".')

add_heading(doc, '15.3 Filter dan Pencarian Referensi', level=2)
add_bullet(doc, 'Pencarian teks: kode, nama, atau kategori')
add_bullet(doc, 'Filter status: Semua / Aktif / Tidak Aktif')
add_bullet(doc, 'Paginasi: 20 item per halaman')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 16 — DOKUMEN ASN
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 16 — DOKUMEN ASN', level=1)
add_para(doc,
    'Halaman Dokumen ASN (menu: SIDATA → Dokumen ASN) memungkinkan pengunggahan dan pengunduhan '
    'dokumen kepegawaian yang dilampirkan ke profil ASN tertentu.')

add_heading(doc, '16.1 Format File yang Didukung', level=2)
add_table_with_header(doc,
    ['Format', 'Ekstensi', 'Ukuran Maksimum'],
    [
        ['PDF', '.pdf', '10 MB'],
        ['Gambar JPEG', '.jpg / .jpeg', '10 MB'],
        ['Gambar PNG', '.png', '10 MB'],
        ['Word Document', '.docx', '10 MB'],
        ['Excel Spreadsheet', '.xlsx', '10 MB'],
    ],
    col_widths=[4.5, 4.0, 4.0]
)

add_heading(doc, '16.2 Cara Mengunggah Dokumen', level=2)
add_numbered(doc, 'Cari ASN yang dituju menggunakan kotak pencarian (NIP, nama, jabatan).')
add_numbered(doc, 'Klik kartu ASN untuk memilihnya. Kartu yang dipilih akan tampil dengan latar berwarna.')
add_numbered(doc, 'Isi kolom "Jenis Dokumen" (contoh: SK Pengangkatan, Ijazah, Surat Keterangan).')
add_numbered(doc, 'Klik "Pilih File" dan pilih file yang akan diunggah.')
add_numbered(doc, 'Klik "Upload". File akan diunggah dan muncul di tabel dokumen di bawah.')

add_heading(doc, '16.3 Mengunduh dan Menghapus Dokumen', level=2)
add_bullet(doc, 'Unduh: Klik tombol "Unduh" pada baris dokumen untuk mengunduh file.')
add_bullet(doc, 'Hapus: Klik tombol "Hapus" untuk menonaktifkan dokumen dari profil ASN.')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 17 — LAPORAN
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 17 — LAPORAN SIDATA', level=1)
add_para(doc,
    'Halaman Laporan (menu: SIDATA → Laporan) menyajikan ringkasan eksekutif kondisi data ASN '
    'secara komprehensif, ditujukan untuk pimpinan dan pengambil keputusan.')

add_heading(doc, '17.1 Kartu KPI Laporan', level=2)
add_table_with_header(doc,
    ['KPI', 'Deskripsi'],
    [
        ['Total ASN', 'Jumlah total ASN dalam database'],
        ['Unit Organisasi', 'Jumlah unit kerja yang terdaftar'],
        ['Total Jabatan', 'Jumlah referensi jabatan aktif'],
        ['Jenis Jabatan', 'Jumlah jenis jabatan (Struktural/Fungsional/Pelaksana)'],
        ['Total Batch Import', 'Jumlah seluruh batch yang pernah dibuat'],
        ['Batch Committed', 'Jumlah batch yang sudah berhasil dicommit'],
        ['Batch Bermasalah', 'Jumlah batch dengan status FAILED atau bermasalah'],
        ['Quality Score', 'Persentase kualitas data keseluruhan'],
    ],
    col_widths=[4.5, 12.0]
)

add_heading(doc, '17.2 Narasi Eksekutif', level=2)
add_para(doc,
    'Bagian "Narasi Eksekutif" secara otomatis menghasilkan teks ringkasan dalam bahasa Indonesia '
    'berdasarkan data terkini. Narasi mencakup:')
add_bullet(doc, 'Ringkasan Kondisi Data: total ASN, unit kerja, referensi jabatan, dan jumlah batch')
add_bullet(doc, 'Catatan Tindak Lanjut: skor kualitas, baris bermasalah, status batch')
add_bullet(doc, 'Rekomendasi: tindakan yang disarankan berdasarkan kondisi data')

add_heading(doc, '17.3 Integrasi RHK SIDATA', level=2)
add_para(doc,
    'Bagian bawah halaman laporan menampilkan panel Laporan RHK yang terhubung dengan '
    'sistem Kinerja Bidang:')
add_bullet(doc, 'RHK 5 (SIK): Laporan sinkronisasi data SIASN')
add_bullet(doc, 'RHK 6 (DAT): Laporan kualitas data ASN')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 18 — KONSEP DAN TERMINOLOGI
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 18 — KONSEP DAN TERMINOLOGI', level=1)

add_table_with_header(doc,
    ['Istilah', 'Definisi'],
    [
        ['ASN', 'Aparatur Sipil Negara — pegawai pemerintah (PNS atau PPPK)'],
        ['SIASN', 'Sistem Informasi ASN — sistem BKN Pusat yang menjadi sumber data utama'],
        ['SIDATA', 'Modul sistem data ASN dalam platform SILAKAP yang dikelola BKPSDM lokal'],
        ['Batch', 'Satu unggahan file yang memproses sekumpulan baris data sekaligus'],
        ['Staging Area', 'Area penyimpanan sementara data sebelum diverifikasi dan dicommit ke master'],
        ['Master Data', 'Database resmi dan final yang digunakan oleh seluruh sistem SILAKAP'],
        ['Commit', 'Operasi final yang memindahkan data dari staging ke master, bersifat permanen'],
        ['Mapping', 'Proses mencocokkan nilai teks (nama jabatan, unit kerja) dengan ID referensi master'],
        ['BUP', 'Batasan Usia Pensiun — usia maksimal pegawai sesuai peraturan yang berlaku'],
        ['TMT', 'Terhitung Mulai Tanggal — tanggal efektif berlakunya suatu status/jabatan/pensiun'],
        ['Quality Score', 'Persentase ASN yang memiliki semua field kritis terisi lengkap'],
        ['Invalid Row', 'Baris data yang gagal validasi dan tidak dapat dicommit'],
        ['Warning Row', 'Baris data yang lolos validasi tapi memiliki catatan kualitas'],
        ['Needs Review', 'Baris data yang memerlukan konfirmasi manual sebelum dicommit'],
        ['Unmapped Row', 'Baris yang nilai teksnya tidak berhasil dipetakan ke referensi master'],
        ['NIP', 'Nomor Induk Pegawai — identitas unik ASN (18 digit)'],
        ['NIK', 'Nomor Induk Kependudukan — nomor KTP 16 digit'],
        ['Golongan', 'Tingkatan kepangkatan ASN (I/a hingga IV/e)'],
        ['Eselon', 'Tingkatan jabatan struktural (Eselon I hingga IV)'],
        ['PPPK', 'Pegawai Pemerintah dengan Perjanjian Kerja — ASN non-PNS'],
        ['RBAC', 'Role-Based Access Control — sistem hak akses berbasis peran'],
        ['Pipeline', 'Rangkaian langkah berurutan dalam proses pemutakhiran data'],
        ['Rekonsiliasi', 'Proses membandingkan dua sumber data untuk menemukan perbedaan'],
        ['Audit Trail', 'Catatan permanen semua operasi yang dilakukan di sistem'],
    ],
    col_widths=[4.0, 12.5]
)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 19 — FAQ
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 19 — PERTANYAAN UMUM (FAQ)', level=1)

faqs = [
    ('Mengapa ada ASN yang tidak muncul di direktori?',
     'Kemungkinan: (1) data belum diimport dari SIASN, (2) batch belum dicommit, atau '
     '(3) ASN berstatus nonaktif dan filter sedang diset ke "Aktif". Coba hapus semua filter.'),
    ('Berapa lama proses upload dan mapping satu batch?',
     'Bergantung pada ukuran file. Batch kecil (<500 baris) biasanya selesai dalam 1–2 menit. '
     'Batch besar (>2000 baris) dapat memakan waktu 5–10 menit.'),
    ('Bisakah saya membatalkan commit yang sudah dilakukan?',
     'Tidak. Commit bersifat final dan tidak dapat dirollback. Jika ada kesalahan, '
     'unggah batch koreksi dan commit ulang untuk menimpa data yang salah.'),
    ('Apa yang harus dilakukan jika banyak baris berstatus UNMAPPED?',
     'Kemungkinan referensi master belum diperbarui. Lakukan Import Referensi terlebih dahulu '
     'dengan data jabatan/golongan/unit terbaru dari BKN, lalu jalankan ulang Auto-Map.'),
    ('Mengapa Quality Score turun setelah commit batch baru?',
     'Batch baru mungkin mengandung ASN dengan field kosong yang sebelumnya tidak ada di master. '
     'Tinjau kartu "Tanpa Unit Kerja", "Tanpa Jabatan", dst. di dashboard untuk detail.'),
    ('Bagaimana cara melihat siapa yang melakukan commit terakhir?',
     'Buka menu Log Sinkronisasi, filter berdasarkan aksi "COMMIT_ASN" atau "COMMIT_REFERENCE". '
     'Setiap entri menampilkan nama pengguna dan timestamp operasi.'),
    ('Apakah data SIDATA otomatis tersinkron dengan BKN SIASN?',
     'Tidak otomatis. Sinkronisasi dilakukan secara manual melalui pipeline Import SIASN. '
     'Disarankan melakukan pemutakhiran minimal satu kali per semester atau setelah ada perubahan besar.'),
    ('Apa perbedaan Import SIASN dan Import Excel?',
     'Import SIASN menggunakan format standar file export dari portal BKN. '
     'Import Excel (saat ini placeholder) akan mendukung format Excel internal BKPSDM. '
     'Gunakan Import SIASN untuk data resmi BKN.'),
    ('Bagaimana cara mengetahui bahwa data jabatan ASN sudah benar?',
     'Gunakan fitur Rekonsiliasi untuk membandingkan batch terbaru dengan master. '
     'Perhatikan baris berjenis "BEDA_DATA" untuk jabatan yang berubah.'),
    ('Apakah pengeditan via Tab Edit di Profil ASN dicatat di log?',
     'Ya, setiap perubahan field dicatat dengan timestamp dan identitas pengguna. '
     'Namun, edit manual tidak membuat batch sehingga tidak muncul di Riwayat Import.'),
]

for i, (q, a) in enumerate(faqs):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f'Q{i+1}: {q}')
    r.font.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x0D, 0x47, 0x2E)
    add_para(doc, f'A: {a}', size=10)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# BAB 20 — LAMPIRAN: ALUR KERJA LENGKAP
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'BAB 20 — LAMPIRAN: ALUR KERJA LENGKAP', level=1)

add_heading(doc, 'Lampiran A — Alur Pemutakhiran Data ASN Rutin', level=2)
add_info_box(doc, 'Gunakan alur ini setiap siklus pemutakhiran data dari BKN SIASN', [
    '',
    '  MULAI',
    '    ↓',
    '  [Operator] Ekspor data ASN dari portal BKN SIASN',
    '    ↓',
    '  [Operator] Upload file ke SIDATA (Import SIASN)',
    '    ↓',
    '  [Reviewer] Jalankan Auto-Map (Mapping Referensi)',
    '    ↓',
    '  Ada baris UNMAPPED atau NEEDS_REVIEW?',
    '    ├── Ya → [Reviewer] Selesaikan mapping manual',
    '    └── Tidak ↓',
    '  [Reviewer] Tinjau Rekonsiliasi (BEDA_DATA, HANYA_DI_BATCH, HANYA_DI_MASTER)',
    '    ↓',
    '  Ada selisih perlu konfirmasi?',
    '    ├── Ya → [Kabid/Admin] Konfirmasi atau tolak perubahan',
    '    └── Tidak ↓',
    '  [Admin] Commit Batch ke Master',
    '    ↓',
    '  [Admin] Verifikasi di Dashboard (Quality Score naik, batch berstatus COMMITTED)',
    '    ↓',
    '  [Admin] Periksa Log Sinkronisasi',
    '    ↓',
    '  SELESAI',
], 'F1F8E9')

add_heading(doc, 'Lampiran B — Alur Import Referensi Baru', level=2)
add_info_box(doc, 'Gunakan saat ada jabatan/golongan/unit baru yang belum ada di master', [
    '',
    '  MULAI',
    '    ↓',
    '  [Admin] Siapkan file Excel referensi (kode, nama)',
    '    ↓',
    '  [Admin] Upload via Import Referensi (pilih jenis referensi)',
    '    ↓',
    '  Sistem validasi format file',
    '    ├── Invalid → Perbaiki file, upload ulang',
    '    └── Valid ↓',
    '  [Admin] Klik Commit pada batch referensi',
    '    ↓',
    '  Referensi baru tersedia di master dan dapat digunakan untuk mapping ASN',
    '    ↓',
    '  SELESAI',
], 'E3F2FD')

add_heading(doc, 'Lampiran C — Matriks Tanggung Jawab per Langkah Pipeline', level=2)
add_table_with_header(doc,
    ['Langkah', 'Operator Import', 'Reviewer Mapping', 'Kabid/Analis Madya', 'Admin BKPSDM'],
    [
        ['Upload SIASN', 'EKSEKUSI', 'PANTAU', 'PANTAU', 'EKSEKUSI'],
        ['Auto-Map', '—', 'EKSEKUSI', 'TINJAU', 'EKSEKUSI'],
        ['Resolve Unmapped', '—', 'EKSEKUSI', 'PANTAU', 'EKSEKUSI'],
        ['Rekonsiliasi', '—', 'EKSEKUSI', 'KONFIRMASI', 'EKSEKUSI'],
        ['Commit ASN', '—', '—', '—', 'EKSEKUSI'],
        ['Verifikasi Log', 'PANTAU', 'PANTAU', 'PANTAU', 'EKSEKUSI'],
    ],
    col_widths=[4.0, 3.5, 3.5, 3.5, 2.0]
)

add_heading(doc, 'Lampiran D — Checklist Pemutakhiran Data Periodik', level=2)
add_para(doc, 'Gunakan checklist ini setiap siklus pemutakhiran:')
checklists = [
    '[ ] File SIASN terbaru sudah diunduh dari portal BKN',
    '[ ] Import SIASN berhasil (status batch: VALIDATED)',
    '[ ] Auto-Map sudah dijalankan (status: MAPPED)',
    '[ ] Tidak ada baris UNMAPPED yang tersisa',
    '[ ] Semua baris NEEDS_REVIEW sudah dikonfirmasi',
    '[ ] Rekonsiliasi sudah ditinjau dan selisih sudah dikonfirmasi',
    '[ ] Commit sudah dilakukan oleh Admin (status: COMMITTED)',
    '[ ] Dashboard Quality Score sudah diperiksa',
    '[ ] Log Sinkronisasi sudah diverifikasi',
    '[ ] Laporan eksekutif sudah dicetak/dikirim ke pimpinan (jika diperlukan)',
]
for item in checklists:
    add_bullet(doc, item)

add_divider(doc)

# Footer note
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
r = p.add_run('Dokumen ini adalah panduan resmi penggunaan modul SIDATA dalam platform SILAKAP V1.0. '
    'Untuk pertanyaan teknis, hubungi administrator sistem BKPSDM.')
r.font.size = Pt(9)
r.font.italic = True
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
r = p.add_run(f'© {datetime.date.today().year} BKPSDM — SILAKAP V1.0 — Versi Dokumen 1.0 — Dicetak: {datetime.date.today().strftime("%d %B %Y")}')
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ─── Save ──────────────────────────────────────────────────────────────────────

output_path = r'd:\Silakap_V1.0\SIDATA_Panduan_Penggunaan.docx'
doc.save(output_path)
print(f'[OK] Dokumen berhasil dibuat: {output_path}')
